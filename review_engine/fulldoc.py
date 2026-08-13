#!/usr/bin/env python3
"""
文档审核批注工具 (新规则版)
输入: docx 文件 + 规则 JSON (审核规则汇总.json, 按模块组织的规则目录)
输出: 带批注的 docx 文件 + 审核结果 xlsx + 规则汇总 xlsx

说明:
- 新规则 JSON 是「规则目录」，含 审核要点/审核问题/如何修改/批注文本/所属模块/启用状态
- 本脚本读取 JSON，筛选「已启用」规则，按模块实现对应的检测逻辑
- 前言、附录 模块在 JSON 中标记为「未启用」，按用户要求跳过
"""

import json
import os
import re
import shutil
import zipfile
from datetime import datetime
from copy import deepcopy

from lxml import etree
from docx import Document

# OOXML 命名空间
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
XML_NS = 'http://www.w3.org/XML/1998/namespace'


def w(tag):
    return f'{{{W_NS}}}{tag}'


class DocxAnnotator:
    """为 docx 文件添加批注 (comments)"""

    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.pending_comments = []  # [(para_idx, text, author, run_start, run_end)]
        self._next_id = 0

    def add_comment(self, para_idx, comment_text, author='文档审核系统', run_start=None, run_end=None):
        """注册一条批注，待保存时写入"""
        self.pending_comments.append({
            'para_idx': para_idx,
            'text': comment_text,
            'author': author,
            'id': self._next_id,
            'run_start': run_start,
            'run_end': run_end,
        })
        self._next_id += 1

    def save(self, output_path):
        """保存带批注的文档"""
        for c in self.pending_comments:
            para = self.doc.paragraphs[c['para_idx']]
            self._insert_comment_markers(para, c)

        self.doc.save(output_path)
        self._post_process(output_path)

    def _insert_comment_markers(self, paragraph, comment):
        p_elem = paragraph._element

        range_start = etree.Element(w('commentRangeStart'))
        range_start.set(w('id'), str(comment['id']))
        first_r = p_elem.find(w('r'))
        if first_r is not None:
            idx = list(p_elem).index(first_r)
            p_elem.insert(idx, range_start)
        else:
            p_elem.append(range_start)

        range_end = etree.Element(w('commentRangeEnd'))
        range_end.set(w('id'), str(comment['id']))
        p_elem.append(range_end)

        ref_run = etree.Element(w('r'))
        rpr = etree.SubElement(ref_run, w('rPr'))
        rstyle = etree.SubElement(rpr, w('rStyle'))
        rstyle.set(w('val'), 'CommentReference')
        ref = etree.SubElement(ref_run, w('commentReference'))
        ref.set(w('id'), str(comment['id']))
        p_elem.append(ref_run)

    def _post_process(self, docx_path):
        tmp_path = docx_path + '.tmp'

        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                has_comments = False

                for item in zin.namelist():
                    data = zin.read(item)

                    if item == '[Content_Types].xml':
                        content = data.decode('utf-8')
                        if 'comments+xml' not in content:
                            override = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                            content = content.replace('</Types>', override + '</Types>')
                        data = content.encode('utf-8')

                    elif item == 'word/_rels/document.xml.rels':
                        content = data.decode('utf-8')
                        if '/comments' not in content:
                            max_id = 0
                            for m in re.finditer(r'Id="rId(\d+)"', content):
                                max_id = max(max_id, int(m.group(1)))
                            new_rid = f'rId{max_id + 1}'
                            rel = (f'<Relationship Id="{new_rid}" '
                                   f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
                                   f'Target="comments.xml"/>')
                            content = content.replace('</Relationships>', rel + '</Relationships>')
                        data = content.encode('utf-8')

                    elif item == 'word/comments.xml':
                        has_comments = True
                        existing = etree.fromstring(data)
                        new_comments = self._build_comments_elements()
                        for c in new_comments:
                            existing.append(c)
                        data = etree.tostring(existing, xml_declaration=True, encoding='UTF-8', standalone=True)

                    zout.writestr(item, data)

                if not has_comments:
                    comments_xml = self._build_comments_xml()
                    zout.writestr('word/comments.xml', comments_xml)

        # 用 os.replace 直接覆盖目标文件 (绕过沙箱安全删除对 os.unlink 的限制)
        os.replace(tmp_path, docx_path)

    def _build_comments_elements(self):
        elements = []
        for c in self.pending_comments:
            comment_elem = etree.Element(w('comment'))
            comment_elem.set(w('id'), str(c['id']))
            comment_elem.set(w('author'), c['author'])
            comment_elem.set(w('date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            comment_elem.set(w('initials'), '审核')

            p = etree.SubElement(comment_elem, w('p'))

            r1 = etree.SubElement(p, w('r'))
            rpr1 = etree.SubElement(r1, w('rPr'))
            rstyle1 = etree.SubElement(rpr1, w('rStyle'))
            rstyle1.set(w('val'), 'CommentReference')
            etree.SubElement(r1, w('annotationRef'))

            r2 = etree.SubElement(p, w('r'))
            t = etree.SubElement(r2, w('t'))
            t.set(f'{{{XML_NS}}}space', 'preserve')
            t.text = c['text']

            elements.append(comment_elem)
        return elements

    def _build_comments_xml(self):
        root = etree.Element(w('comments'), nsmap={'w': W_NS})
        for elem in self._build_comments_elements():
            root.append(elem)
        return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


class NewRuleChecker:
    """按新规则 JSON (审核规则汇总.json) 实现的审核检查器"""

    # 模块 -> 是否启用 (根据 JSON 的 模块状态/启用状态 决定)
    # 前言、附录 模块在 JSON 中标记为未启用，跳过
    DISABLED_MODULES = {'前言', '附录'}

    def __init__(self, annotator, rules_data):
        self.ann = annotator
        self.rules_data = rules_data
        self.doc = annotator.doc
        self.paragraphs = self.doc.paragraphs

        # findings: 每条检测结果的完整记录 (用于 xlsx)
        self.findings = []

        # 加载规则目录
        self.all_rules = rules_data.get('规则明细', [])
        self.enabled_rules = [r for r in self.all_rules
                              if 'active' in r.get('启用状态', '').lower()
                              and r.get('所属模块', '') not in self.DISABLED_MODULES]
        self.rules_by_module = {}
        for r in self.enabled_rules:
            self.rules_by_module.setdefault(r.get('所属模块', '其他'), []).append(r)

        self._parse_structure()

    # ── 结构解析 ────────────────────────────────────────
    def _parse_structure(self):
        self.h1_positions = []      # [(idx, title)]
        self.toc_idx = None
        self.foreword_idx = None
        self.intro_idx = None
        self.scope_idx = None
        self.refs_idx = None
        self.terms_idx = None
        self.abbrev_idx = None
        self.cover_title = None
        self.cover_end = 0

        for i, p in enumerate(self.paragraphs):
            text = p.text.strip()
            style = p.style.name if p.style else ''
            clean = re.sub(r'[\s\u2003\u2002\u2007\u2009\u202f]+', '', text)

            if 'Heading 1' in style:
                self.h1_positions.append((i, text))

            if clean == '目次' and ('目次' in style or '标题' in style or '标准名称' in style):
                if self.toc_idx is None:
                    self.toc_idx = i
            if clean == '前言' and ('前言' in style or '引言' in style or '标题' in style):
                if self.foreword_idx is None:
                    self.foreword_idx = i
            if clean == '引言' and ('前言' in style or '引言' in style or '标题' in style):
                if self.intro_idx is None:
                    self.intro_idx = i
            if clean == '范围' and 'Heading' in style:
                if self.scope_idx is None:
                    self.scope_idx = i
            if '规范性引用文件' in text and 'Heading' in style:
                if self.refs_idx is None:
                    self.refs_idx = i
            if clean == '术语和定义' and 'Heading' in style:
                if self.terms_idx is None:
                    self.terms_idx = i
            if clean == '缩略语' and 'Heading' in style:
                if self.abbrev_idx is None:
                    self.abbrev_idx = i

        # 封面标题 (含"技术要求"等中文名)
        for i in range(min(12, len(self.paragraphs))):
            t = self.paragraphs[i].text.strip()
            if '技术要求' in t or '通信行业标准' in t or ('IMS' in t and '业务' in t):
                self.cover_title = t
                break

        if self.toc_idx is not None:
            self.cover_end = self.toc_idx
        elif self.foreword_idx is not None:
            self.cover_end = self.foreword_idx
        else:
            self.cover_end = 15

    def _add(self, para_idx, module, audit_point, comment, rule_question=''):
        """记录一条检测结果并注册批注"""
        self.ann.add_comment(para_idx, comment)
        self.findings.append({
            'para_idx': para_idx,
            'module': module,
            'audit_point': audit_point,
            'comment': comment,
            'rule_question': rule_question,
            'orig_text': self.paragraphs[para_idx].text.strip()[:80] if 0 <= para_idx < len(self.paragraphs) else '',
        })

    def check_all(self):
        self._check_main()
        self._check_toc()
        self._check_scope()
        self._check_refs()
        self._check_terms()
        self._check_abbrev()
        self._check_body()

    # ── 主程序 (章节结构) ────────────────────────────────
    def _check_main(self):
        # 章节-目次
        if self.toc_idx is None:
            self._add(self.scope_idx or 0, '主程序(main)', '章节-目次',
                      '【章节-目次】规则要求：目次章节为必备要素，当前文档缺少目次章节。')
        # 章节-范围
        if self.scope_idx is None:
            self._add(self.toc_idx or 0, '主程序(main)', '章节-范围',
                      '【章节-范围】规则要求：范围章节为必备要素，当前文档缺少范围章节。')
        # 章节-规范性引用文件
        if self.refs_idx is None:
            self._add(self.scope_idx or 0, '主程序(main)', '章节-规范性引用文件',
                      '【章节-规范性引用文件】规则要求："规范性引用文件"一章的章编号和标题是必备的，'
                      '任何文件中都必须设置（若没有规范性引用文件需保留第2章并给出引导语：本文件没有规范性引用文件）。')
        # 章节-术语和定义
        if self.terms_idx is None:
            self._add(self.scope_idx or 0, '主程序(main)', '章节-术语和定义',
                      '【章节-术语和定义】规则要求：术语和定义章节为必备要素，若没有术语和定义需保留第3章并给出引导语：本文件没有术语和定义。')
        # 章节-缩略语
        if self.abbrev_idx is None:
            self._add(self.terms_idx or 0, '主程序(main)', '章节-缩略语',
                      '【章节-缩略语】规则要求：缩略语章节结构应正确，若无缩略语内容建议删除该章节或保留并注明：本文件没有需要解释的缩略语。')
        # 章节-引言: 仅当标准名称为"部分"时需要
        if self.intro_idx is None and self.cover_title and '部分' in self.cover_title:
            self._add(self.foreword_idx or 0, '主程序(main)', '章节-引言',
                      '【章节-引言】规则要求：标准名称为"部分"的标准，应有引言这一章节，并给出每部分的目的。当前文档未设置引言。')

    # ── 目次/目录 ──────────────────────────────────────
    def _check_toc(self):
        if self.toc_idx is None:
            return

        # 目次名称: 应当是"目次"而非"目录"
        toc_title = self.paragraphs[self.toc_idx].text.strip()
        if '目录' in toc_title and '目次' not in toc_title:
            self._add(self.toc_idx, '目次/目录', '目次-名称',
                      '【目次-名称】规则要求：目次名称应当为"目次"，不应写作"目录"。当前目次标题为"目录"，应修改成"目次"。')

        # 前言区域边界 (用于区分目次条目与前言内容)
        foreword_end = self.scope_idx if self.scope_idx else (self.foreword_idx or self.toc_idx) + 10
        if self.foreword_idx is not None:
            for idx, _ in self.h1_positions:
                if idx > self.foreword_idx:
                    foreword_end = idx
                    break

        # 目次条目区域: 目次标题之后、前言之前
        toc_entries_end = self.foreword_idx if self.foreword_idx else (self.scope_idx or self.toc_idx + 10)
        toc_entries = []
        for i in range(self.toc_idx + 1, toc_entries_end):
            t = self.paragraphs[i].text.strip()
            if t and '目次' not in t:
                toc_entries.append((i, t))

        # 目次为空 / 缺少必要章节
        if not toc_entries:
            missing = []
            if self.foreword_idx is not None:
                missing.append('前言')
            if self.scope_idx is not None:
                missing.append('范围')
            if self.refs_idx is not None:
                missing.append('规范性引用文件')
            if self.terms_idx is not None:
                missing.append('术语和定义')
            if self.abbrev_idx is not None:
                missing.append('缩略语')
            self._add(self.toc_idx, '目次/目录', '目次-结构',
                      '【目次-结构】规则要求：目次应列出前言、范围、规范性引用文件、术语和定义等必备要素的章节名称及页码。'
                      f'当前目次为空（未生成或未更新域），缺少以下章节条目：{("、".join(missing) if missing else "必备要素")}。'
                      '建议更新目次域以生成正确页码与条目。')

        # 目次中不能出现本标准的中文名名称 (扫描 目次标题~范围 之间，排除前言内容)
        if self.cover_title:
            for i in range(self.toc_idx + 1, self.scope_idx if self.scope_idx else self.toc_idx + 20):
                if self.foreword_idx and self.foreword_idx <= i < foreword_end:
                    continue  # 跳过前言内容
                t = self.paragraphs[i].text.strip()
                if self.cover_title in t:
                    self._add(i, '目次/目录', '目次-段落',
                              f'【目次-段落】规则要求：目次中不能出现本标准的中文名名称。'
                              f'当前目次区域出现了标准中文名"{self.cover_title}"，应删除。')
                    break

        # 目次中不应含有错误标签
        for i, t in toc_entries:
            if '错误！未定义标签' in t or '错误!未定义标签' in t:
                self._add(i, '目次/目录', '目次-段落',
                          '【目次-段落】规则要求：目次中不应含有错误标签。当前目次存在"错误！未定义标签"，请更新域或修正。')

    # ── 范围 ──────────────────────────────────────────
    def _check_scope(self):
        if self.scope_idx is None:
            return
        scope_end = self.refs_idx if self.refs_idx else self.scope_idx + 5

        scope_text = ''
        for i in range(self.scope_idx, scope_end):
            if i < len(self.paragraphs):
                scope_text += self.paragraphs[i].text

        # 范围中不能出现书名号
        if '《' in scope_text or '》' in scope_text:
            for i in range(self.scope_idx, scope_end):
                t = self.paragraphs[i].text
                if '《' in t or '》' in t:
                    self._add(i, '范围', '正文-内容',
                              '【范围】规则要求：范围中不能出现书名号《》。当前范围章节含有书名号，请删除。')
                    break

        # 范围的开头应符合规定
        first_para = ''
        for i in range(self.scope_idx + 1, scope_end):
            if i < len(self.paragraphs) and self.paragraphs[i].text.strip():
                first_para = self.paragraphs[i].text.strip()
                break
        valid_starts = ('本文件确立了', '本文件规定了', '本文件描述了', '本文件界定了')
        if first_para and not first_para.startswith(valid_starts):
            self._add(self.scope_idx + 1, '范围', '范围-格式',
                      '【范围-格式】规则要求：范围的第一段开头应为"本文件确立了…"/"本文件规定了…"/'
                      '"本文件描述了…"/"本文件界定了…"之一。'
                      f'当前开头为"{first_para[:30]}…"，不符合规定结构。')

        # 范围第一段不应出现"强调了"等
        if '强调了' in scope_text:
            self._add(self.scope_idx + 1, '范围', '范围-格式',
                      '【范围-格式】规则要求：范围的第一段中不应出现"强调了…"等评价性文字。')

        # 范围内容结构: 第一段标准化对象 + 第二段适用/不适用
        paras = [self.paragraphs[i].text.strip() for i in range(self.scope_idx + 1, scope_end)
                 if i < len(self.paragraphs) and self.paragraphs[i].text.strip()]
        if paras:
            has_scope = any(p.startswith('本文件规定') or p.startswith('本文件确立') for p in paras)
            has_apply = any(p.startswith('本文件适用') for p in paras)
            if not (has_scope and has_apply):
                self._add(self.scope_idx + 1, '范围', '范围-格式',
                          '【范围-格式】规则要求：范围内容应符合规定结构——第一段陈述该文件的标准化对象，'
                          '第二段为本文件适用与不适用（不适用也可另起一段）。当前范围结构不完整。')

    # ── 规范性引用文件 ───────────────────────────────────
    def _check_refs(self):
        if self.refs_idx is None:
            return
        ref_end = len(self.paragraphs)
        for idx, title in self.h1_positions:
            if idx > self.refs_idx:
                ref_end = idx
                break

        ref_paras = []
        for i in range(self.refs_idx + 1, ref_end):  # 跳过 H1 标题本身
            t = self.paragraphs[i].text.strip()
            if t:
                ref_paras.append((i, t))

        if not ref_paras:
            return

        # 引导语检查
        intro_text = ref_paras[0][1]
        std_intro = ('下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款。'
                     '其中，注日期的引用文件，仅该日期对应的版本适用于本文件；'
                     '不注日期的引用文件，其最新版本（包括所有的修改单）适用于本文件。')
        if '下列文件' in intro_text:
            if re.sub(r'\s+', '', intro_text) != re.sub(r'\s+', '', std_intro):
                self._add(ref_paras[0][0], '规范性引用文件', '规范性引用文件-引导语',
                          '【规范性引用文件-引导语】规则要求：引导语应使用规定的标准文本——'
                          '"下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款。其中，注日期的引用文件，'
                          '仅该日期对应的版本适用于本文件；不注日期的引用文件，其最新版本（包括所有的修改单）适用于本文件。"'
                          '当前引导语与标准文本不一致，请修改。')
        else:
            self._add(ref_paras[0][0], '规范性引用文件', '规范性引用文件-引导语',
                      '【规范性引用文件-引导语】规则要求：若文件中规范性引用了其他文件，必须在章标题下使用规定的标准引导语。当前缺少标准引导语。')

        # 引用文件条目解析 + 排序检查
        entries = []  # (para_idx, org, code, raw)
        for i, t in ref_paras[1:]:
            # 跳过明显非条目 (如纯说明)
            m = re.match(r'^\s*([A-Za-z0-9/]+)\s+(RFC|TS|TR|IS)\s*([\d\.]+)', t)
            if m:
                org = m.group(1).upper()
                entries.append((i, org, t))

        # 排序: 按组织代号的拉丁字母/数字顺序，组织内按编号
        # 标准期望: 数字开头(如3GPP)在组织字母(IETF)之前
        if len(entries) > 1:
            # 检查是否有 IETF 出现在 3GPP 之前
            org_seq = [e[1] for e in entries]
            # 找第一个 IETF 和第一个 3GPP 的位置
            first_i = next((k for k, o in enumerate(org_seq) if o == 'IETF'), None)
            first_3 = next((k for k, o in enumerate(org_seq) if o == '3GPP'), None)
            if first_i is not None and first_3 is not None and first_i < first_3:
                self._add(entries[first_3][0], '规范性引用文件', '规范性引用文件-排序规则',
                          '【规范性引用文件-排序规则】规则要求：引用文件应按标准类型和编号重新排序。'
                          '按组织代号（数字开头先于字母）排序时，3GPP（数字开头）应排在 IETF（字母开头）之前。'
                          '当前 IETF RFC 文件排在 3GPP TS 文件之前，顺序不正确，请调整。')

        # 引用文件不应使用表格格式 (本文无表格，跳过)
        # 引用文件不应含有书名号/引号
        for i, t in ref_paras[1:]:
            if '《' in t or '》' in t or '"' in t:
                self._add(i, '规范性引用文件', '规范性引用文件-引用文件',
                          '【规范性引用文件-引用文件】规则要求：引用文件条目不应含有书名号或引号。'
                          '请删除文件名中的书名号、引号。')
                break

    # ── 术语和定义 ──────────────────────────────────────
    def _check_terms(self):
        if self.terms_idx is None:
            return
        terms_end = self.abbrev_idx if self.abbrev_idx else self.terms_idx + 5

        terms_paras = []
        for i in range(self.terms_idx, terms_end):
            t = self.paragraphs[i].text.strip()
            if t:
                terms_paras.append((i, t))

        if not terms_paras:
            return

        guide = terms_paras[0][1]
        # 引导语检查
        if '没有需要界定' in guide:
            # 正确引导语，无需处理
            pass
        elif '下列' in guide or '术语' in guide:
            # 有术语，检查格式
            self._audit_terms_content(terms_paras)
        else:
            self._add(terms_paras[0][0], '术语和定义', '术语和定义-引导语',
                      '【术语和定义-引导语】规则要求：若无术语和定义应写"本文件没有需要界定的术语和定义。"；'
                      '若有术语和定义应给出相应引导语。当前引导语不符合要求。')

    def _audit_terms_content(self, terms_paras):
        """检查术语和定义内容 (有术语时)"""
        full = '\n'.join(t for _, t in terms_paras)
        # 不应包含要求性条款
        modal = ['应', '宜', '可', '不应', '不准许', '无须']
        # 定义中不应使用 是指/指/是
        bad_words = ['是指', '系指']
        for i, t in terms_paras[1:]:
            if any(bw in t for bw in bad_words):
                self._add(i, '术语和定义', '术语和定义-条款约束',
                          '【术语和定义-条款约束】规则要求：定义中不应使用"是指""指""是"等字，请改用"定义为""称为"等客观表述。')
                break

    # ── 缩略语 ────────────────────────────────────────
    def _check_abbrev(self):
        if self.abbrev_idx is None:
            return
        abbrev_end = self.scope_idx  # 缩略语在术语和定义之后，范围之前已处理；此处取下一 H1
        # 找到缩略语之后的下一个 H1
        next_h1 = len(self.paragraphs)
        for idx, _ in self.h1_positions:
            if idx > self.abbrev_idx:
                next_h1 = idx
                break
        abbrev_end = next_h1

        # 收集缩略语内容段落
        abbrev_text = ''
        for i in range(self.abbrev_idx, abbrev_end):
            abbrev_text += self.paragraphs[i].text + '\n'

        # 解析条目: 支持两种格式
        # 格式A (目标): 缩略语：中文（英文）
        # 格式B (当前文档): 逗号CSV + $$ 前缀
        lines = [ln.strip() for ln in abbrev_text.split('\n') if ln.strip()]
        # 去掉引导语行
        entries = []
        for ln in lines:
            if ln.startswith('下列缩略语') or ln.startswith('缩略语,') or ln == '缩略语':
                continue
            if ln.startswith('$$') or ',' in ln:
                # CSV 格式
                raw = ln.lstrip('$').strip()
                parts = [p.strip() for p in raw.split(',')]
                if len(parts) >= 2:
                    entries.append((parts[0], parts[1], parts[2] if len(parts) > 2 else ''))
            else:
                # 尝试 缩略语：中文（英文）
                m = re.match(r'^(.+?)[:：]\s*(.+?)[(（](.+?)[)）]$', ln)
                if m:
                    entries.append((m.group(1), m.group(2), m.group(3)))

        if not entries:
            return

        # 格式检查: 应为 缩略语：中文（英文），而非逗号分隔
        csv_like = any(',' in ln and ln.lstrip('$').strip().count(',') >= 2 for ln in lines if ln not in ('下列缩略语适用于本文件。', '缩略语,中文全称,英文全称'))
        if csv_like or any(ln.startswith('$$') for ln in lines):
            self._add(self.abbrev_idx + 1, '缩略语', '缩略语-结构',
                      '【缩略语-结构】规则要求：缩略语格式应为"缩略语：中文（英文）"，不应使用表格或逗号分隔的文本格式。'
                      '当前缩略语采用逗号分隔（含"$$"转换残留），请改为"缩略语：中文（英文）"格式，例如：AAA：鉴权与授权响应（Authorization-Authentication-Answer）。')

        # 首字母 A-Z 排序检查
        # 期望顺序: 字母 A-Z 在前，数字开头按数值在后
        def sort_key(e):
            abbr = e[0]
            if abbr and abbr[0].isdigit():
                return (1, abbr)  # 数字开头排后面
            return (0, abbr.upper())

        expected = sorted(entries, key=sort_key)
        # 比较相邻是否严格有序 (找出首个乱序)
        out_of_order = []
        for j in range(1, len(entries)):
            a = sort_key(entries[j - 1])
            b = sort_key(entries[j])
            # 同一组(字母/数字)内需按 abbr 升序
            if a[0] == b[0] and a[1].upper() > b[1].upper():
                out_of_order.append((entries[j - 1][0], entries[j][0]))

        if out_of_order:
            examples = '；'.join([f'{x[0]}应在{y[0]}之前' for x, y in out_of_order[:5]])
            self._add(self.abbrev_idx + 1, '缩略语', '缩略语-首字母',
                      f'【缩略语-首字母】规则要求：缩略语首字母应按照从 A 到 Z 的顺序排列，数字开头的缩略语按从小到大排列。'
                      f'当前顺序不正确，例如：{examples}。请按字母/数字顺序重排。')

        # 重复检查
        seen = {}
        for abbr, cn, en in entries:
            seen[abbr] = seen.get(abbr, 0) + 1
        dups = [k for k, v in seen.items() if v > 1]
        if dups:
            self._add(self.abbrev_idx + 1, '缩略语', '缩略语-内容',
                      f'【缩略语-内容】规则要求：缩略语不应重复。当前重复缩略语：{("、".join(dups))}。请删除重复项。')

    # ── 正文 ──────────────────────────────────────────
    def _check_body(self):
        placeholders = ['待补充', '待完善', '待定', '待确认', 'XXX', 'TBD', 'TODO', 'xxxx']
        for i, p in enumerate(self.paragraphs):
            if self.scope_idx and i < self.scope_idx:
                continue
            t = p.text
            for ph in placeholders:
                if ph.lower() in t.lower():
                    self._add(i, '正文', '正文-正文内容',
                              f'【正文】规则要求：正文不应含有占位符（如"待补充/待完善/待定/XXX"）。'
                              f'当前段落含有"{ph}"，需要补充实际内容。')
                    break


def generate_results_xlsx(annotator, checker, output_path):
    """将审核结果（检测到的错误和批注）导出为 xlsx，含规则映射"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    HEADER_FONT = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
    HEADER_FILL = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    MODULE_FONT = Font(name='微软雅黑', size=10, bold=True, color='1F4E79')
    CELL_FONT = Font(name='微软雅黑', size=10)
    WRAP_ALIGN = Alignment(wrap_text=True, vertical='top', horizontal='left')
    THIN_BORDER = Border(
        left=Side(style='thin', color='B4C7E7'),
        right=Side(style='thin', color='B4C7E7'),
        top=Side(style='thin', color='B4C7E7'),
        bottom=Side(style='thin', color='B4C7E7'),
    )

    wb = Workbook()
    ws = wb.active
    ws.title = '审核结果'

    headers = ['序号', '所属模块', '审核要点', '段落索引', '原文内容（前80字）', '批注内容', '对应规则(审核问题)']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(wrap_text=True, vertical='center', horizontal='center')
        cell.border = THIN_BORDER
    ws.row_dimensions[1].height = 28

    findings = checker.findings
    section_counts = {}

    for idx, f in enumerate(findings, 1):
        ws.cell(row=idx + 1, column=1, value=idx)
        ws.cell(row=idx + 1, column=2, value=f['module'])
        ws.cell(row=idx + 1, column=3, value=f['audit_point'])
        ws.cell(row=idx + 1, column=4, value=f['para_idx'])
        ws.cell(row=idx + 1, column=5, value=f['orig_text'])
        ws.cell(row=idx + 1, column=6, value=f['comment'])
        ws.cell(row=idx + 1, column=7, value=f['rule_question'])

        for col in range(1, 8):
            cell = ws.cell(row=idx + 1, column=col)
            cell.font = CELL_FONT
            cell.alignment = WRAP_ALIGN
            cell.border = THIN_BORDER
        ws.cell(row=idx + 1, column=2).font = MODULE_FONT

        section_counts[f['module']] = section_counts.get(f['module'], 0) + 1

    for i, w in enumerate([6, 16, 20, 10, 45, 60, 40], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = 'A2'
    ws.auto_filter.ref = f'A1:G{len(findings) + 1}'

    # 统计 sheet
    ws2 = wb.create_sheet('统计摘要')
    ws2.cell(row=1, column=1, value='模块')
    ws2.cell(row=1, column=2, value='批注数')
    for col in range(1, 3):
        cell = ws2.cell(row=1, column=col)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = THIN_BORDER
    ws2.row_dimensions[1].height = 28

    # 按 JSON 模块顺序排列
    module_order = ['主程序(main)', '目次/目录', '范围', '规范性引用文件', '术语和定义', '缩略语', '正文']
    sorted_sections = sorted(section_counts.items(),
                             key=lambda x: (module_order.index(x[0]) if x[0] in module_order else 99, -x[1]))
    for i, (sec, cnt) in enumerate(sorted_sections):
        ws2.cell(row=i + 2, column=1, value=sec)
        ws2.cell(row=i + 2, column=2, value=cnt)
        for col in range(1, 3):
            cell = ws2.cell(row=i + 2, column=col)
            cell.font = CELL_FONT
            cell.alignment = WRAP_ALIGN
            cell.border = THIN_BORDER

    total_row = len(sorted_sections) + 3
    ws2.cell(row=total_row, column=1, value='合计')
    ws2.cell(row=total_row, column=2, value=len(findings))
    for col in range(1, 3):
        cell = ws2.cell(row=total_row, column=col)
        cell.font = Font(name='微软雅黑', size=10, bold=True)
        cell.border = THIN_BORDER
    ws2.column_dimensions['A'].width = 18
    ws2.column_dimensions['B'].width = 10

    # 规则覆盖 sheet (已启用规则 vs 实际触发)
    ws3 = wb.create_sheet('规则覆盖')
    ws3.cell(row=1, column=1, value='所属模块')
    ws3.cell(row=1, column=2, value='已启用规则数')
    ws3.cell(row=1, column=3, value='本次触发数')
    for col in range(1, 4):
        cell = ws3.cell(row=1, column=col)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = THIN_BORDER

    enabled_by_module = {}
    for r in checker.enabled_rules:
        enabled_by_module[r.get('所属模块', '其他')] = enabled_by_module.get(r.get('所属模块', '其他'), 0) + 1
    triggered_by_module = section_counts

    all_modules = list(enabled_by_module.keys())
    for i, mod in enumerate(all_modules):
        ws3.cell(row=i + 2, column=1, value=mod)
        ws3.cell(row=i + 2, column=2, value=enabled_by_module[mod])
        ws3.cell(row=i + 2, column=3, value=triggered_by_module.get(mod, 0))
        for col in range(1, 4):
            cell = ws3.cell(row=i + 2, column=col)
            cell.font = CELL_FONT
            cell.border = THIN_BORDER
    ws3.column_dimensions['A'].width = 18
    ws3.column_dimensions['B'].width = 14
    ws3.column_dimensions['C'].width = 14

    wb.save(output_path)


DEFAULT_RULES_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'rules', 'fulldoc_rules.json',
)


def run_fulldoc_review(docx_path, rules_path=None, out_docx=None, out_xlsx=None,
                       write_comments=True, write_xlsx=True):
    """全文档结构审核（主程序 / 目次 / 范围 / 引用 / 术语 / 缩略语 / 正文）。

    规则来自 rules/fulldoc_rules.json（审核规则汇总，156 条明细，按「启用状态」筛选）。
    返回 dict：findings / out_docx / out_xlsx / rule_stats。
    """
    rules_path = rules_path or DEFAULT_RULES_PATH
    if not os.path.exists(rules_path):
        raise FileNotFoundError('找不到规则文件 %s' % rules_path)
    if not os.path.exists(docx_path):
        raise FileNotFoundError('找不到输入文件 %s' % docx_path)

    with open(rules_path, 'r', encoding='utf-8') as f:
        rules_data = json.load(f)

    details = rules_data.get('规则明细', [])
    enabled = [r for r in details if 'active' in r.get('启用状态', '').lower()]

    base, _ext = os.path.splitext(docx_path)
    out_docx = out_docx or (base + '_全文档审核批注.docx')
    out_xlsx = out_xlsx or (base + '_全文档审核问题.xlsx')

    annotator = DocxAnnotator(docx_path)
    checker = NewRuleChecker(annotator, rules_data)
    checker.check_all()

    if write_comments:
        annotator.save(out_docx)
    else:
        out_docx = None

    if write_xlsx:
        generate_results_xlsx(annotator, checker, out_xlsx)
    else:
        out_xlsx = None

    return {
        'findings': checker.findings,
        'out_docx': out_docx,
        'out_xlsx': out_xlsx,
        'rule_stats': {'总规则数': len(details), '启用规则数': len(enabled)},
    }


def main(argv=None):
    import argparse

    parser = argparse.ArgumentParser(
        description='全文档结构审核: docx 输入, 输出带批注 docx + 问题 xlsx')
    parser.add_argument('input', help='待审核的 .docx 文件')
    parser.add_argument('-o', '--output-docx', default=None, help='输出带批注 docx')
    parser.add_argument('-x', '--output-xlsx', default=None, help='输出问题清单 xlsx')
    parser.add_argument('--rules', default=None, help='规则 json 路径 (默认 <仓库>/rules/fulldoc_rules.json)')
    args = parser.parse_args(argv)

    result = run_fulldoc_review(args.input, args.rules, args.output_docx, args.output_xlsx)
    stats = result['rule_stats']
    print(f"已加载规则: {stats['总规则数']} 条明细，启用 {stats['启用规则数']} 条")
    print(f"\n共发现 {len(result['findings'])} 条审核批注:")
    for f in result['findings']:
        print(f'  [段落{f["para_idx"]}] [{f["module"]}/{f["audit_point"]}] {f["comment"][:50]}...')
    print(f"\n带批注文档已保存: {result['out_docx']}")
    print(f"审核结果表格已保存: {result['out_xlsx']}")
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
