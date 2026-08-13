# -*- coding: utf-8 -*-
"""
术语章节审核器
===============
基于《术语审核规则.json》生成的审核代码，用于对标准/规范文档中的
"术语和定义" / "术语、定义和缩略语"章节进行自动化审核。

支持输入：
    - 已解析的章节文本（字符串 / 段落列表）
    - .docx 文件路径（依赖 python-docx）

输出：
    每条规则对应的问题列表，包含位置、问题描述、修改建议。
"""

from __future__ import annotations

import datetime
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable, List, Optional, Tuple

from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.opc.part import XmlPart
from docx.opc.packuri import PackURI


# ---------------------------------------------------------------------------
# 数据模型
# ---------------------------------------------------------------------------

@dataclass
class AuditIssue:
    """单个审核问题。"""
    rule: str                       # 审核要点
    location: str                   # 位置描述
    description: str                # 审核问题
    suggestion: str                 # 如何修改
    severity: str = "error"           # error / warning / info
    details: dict = field(default_factory=dict)
    paragraph_index: int = -1        # 在"章节段落列表"中的索引，用于回写批注；-1 表示不可锚定

    def to_dict(self) -> dict:
        return {
            "审核要点": self.rule,
            "位置": self.location,
            "问题": self.description,
            "修改建议": self.suggestion,
            "级别": self.severity,
            "段落索引": self.paragraph_index,
            **self.details,
        }


@dataclass
class TermEntry:
    """解析出的术语条目。"""
    name_cn: str = ""               # 中文术语
    name_en: str = ""               # 英文术语
    definition: str = ""            # 定义文本
    paragraph_index: int = -1        # 所在段落索引
    raw_text: str = ""              # 原始文本


@dataclass
class TerminologyChapter:
    """术语章节结构。"""
    title: str = ""                 # 章节标题
    guide_sentence: str = ""        # 引导语
    guide_paragraph_index: int = -1 # 引导语段落索引
    sub_sections: List[str] = field(default_factory=list)
    entries: List[TermEntry] = field(default_factory=list)
    body_paragraphs: List[str] = field(default_factory=list)  # 除引导语、子章节标题外的段落
    acronym_section: List[str] = field(default_factory=list)  # "缩略语"子章节正文段落
    declared_abbreviations: set = field(default_factory=set)  # "缩略语"章节已声明的缩写
    source_paragraphs: List[str] = field(default_factory=list)  # 章节段落文本（解析输入，用于回写批注定位）


# ---------------------------------------------------------------------------
# 规则常量
# ---------------------------------------------------------------------------

VALID_GUIDE_PATTERNS = [
    r"^下列术语和定义适用于本文件。$",
    r"^.+界定的术语和定义适用于本文件。$",
    r"^.+界定的以及下列术语和定义适用于本文件。$",
    r"^本文件没有需要界定的术语和定义。$",
]

VALID_GUIDE_PATTERNS_NO_TERMS = [
    r"^本文件没有需要界定的术语和定义。$",
]

REQUIREMENT_WORDS = [
    "应", "宜", "可", "应该", "建议", "可以", "不应", "不准许", "无须",
]

FORBIDDEN_LEADING_WORDS = ["系指", "是指", "指", "是"]

TERMINOLOGY_CHAPTER_TITLES = ["术语和定义", "术语、定义和缩略语"]

# 英文术语后可选的括号说明（半角/全角括号或方头括号），如 (DR)、（OS-P）、【master station】
EN_TERM_BRACKET = r"(\([^()]*\)|（[^（）]*）|【[^【】]*】)?"

# 含"应/可/宜"但本身并非要求性条款的常见复合词。
# 扫描要求性条款前先剔除，避免单字"应/可"误命中"响应/应用/可能/可行"等。
SAFE_COMPOUNDS_WITH_MODAL = [
    # 含"应"
    "响应", "应用", "反应", "相应", "适应", "供应", "对应", "应答", "效应", "应急",
    "应邀", "应付", "应聘", "应允", "理应", "反应",
    # 含"可"
    "可能", "可行", "可信", "可持续", "可见", "可靠", "可取", "认可", "许可",
    "能够", "能力", "可谓", "可贵", "可喜", "可观", "可乘",
    # 含"宜"
    "适宜", "宜人", "便宜",
]


# ---------------------------------------------------------------------------
# 解析器：从段落列表中提取术语章节结构
# ---------------------------------------------------------------------------

class TerminologyParser:
    """将章节段落解析为 TerminologyChapter。"""

    def __init__(self, paragraphs: List[str]):
        self.paragraphs = [p.strip() for p in paragraphs if p.strip()]
        self.sub_sections: List[str] = []
        self.chapter_title: str = ""

    def parse(self, chapter_title: str = "") -> TerminologyChapter:
        """
        解析术语章节。

        参数:
            chapter_title: 章节标题；若为空则自动检测。
        """
        chapter = TerminologyChapter()
        chapter.title = chapter_title or self._detect_title()
        self.chapter_title = chapter.title
        chapter.source_paragraphs = self.paragraphs

        if not self.paragraphs:
            return chapter

        # 第一段通常是引导语；如果第一章节标题在段落中，则跳过
        start_idx = 0
        if self.paragraphs[0] == chapter.title:
            start_idx = 1

        # 查找引导语
        guide_idx, guide_text = self._find_guide_sentence(start_idx)
        chapter.guide_sentence = guide_text
        chapter.guide_paragraph_index = guide_idx

        # 解析子章节（仅针对"术语、定义和缩略语"）
        self.sub_sections = self._extract_sub_sections(guide_idx)
        chapter.sub_sections = self.sub_sections

        # 提取"缩略语"子章节正文与已声明的缩写（供缩略语分号检查排除交叉引用）
        if "缩略语" in self.sub_sections:
            chapter.acronym_section = self._extract_acronym_section(guide_idx)
            chapter.declared_abbreviations = self._extract_declared_abbreviations(
                chapter.acronym_section
            )

        # 若引导语为"无术语需界定"，本章不存在术语条目，不再解析术语条目
        if self._is_no_terms_guide(guide_text):
            chapter.entries = []
            chapter.body_paragraphs = self._extract_body_paragraphs(guide_idx)
            return chapter

        # 解析术语条目
        chapter.entries = self._parse_entries(guide_idx)

        # 保存章节正文段落（用于术语未换行等检查）
        chapter.body_paragraphs = self._extract_body_paragraphs(guide_idx)

        return chapter

    def _is_no_terms_guide(self, guide: str) -> bool:
        """判断引导语是否为'本章无术语需界定'形式。"""
        return bool(re.match(r"^本文件没有需要界定的术语和定义。$", guide or ""))

    def _detect_title(self) -> str:
        if not self.paragraphs:
            return ""
        first = self.paragraphs[0]
        for title in TERMINOLOGY_CHAPTER_TITLES:
            if title in first:
                return title
        return first

    def _find_guide_sentence(self, start_idx: int) -> Tuple[int, str]:
        """查找引导语所在段落。

        优先严格匹配合规引导语；若均不匹配（引导语格式不合规），
        则回退到标题后的首段作为术语边界，仍交由格式/冗余检查去报错，
        避免整章术语因引导语错误而解析失败。
        """
        for i in range(start_idx, min(start_idx + 5, len(self.paragraphs))):
            text = self.paragraphs[i]
            for pattern in VALID_GUIDE_PATTERNS:
                if re.match(pattern, text):
                    return i, text
        # 回退：标题后首段即视为引导语边界
        if start_idx < len(self.paragraphs):
            return start_idx, self.paragraphs[start_idx]
        return -1, ""

    def _extract_sub_sections(self, guide_idx: int) -> List[str]:
        """
        提取直接子章节标题。

        仅在章节标题为"术语、定义和缩略语"时生效，
        且只认可标准子章节名称："术语和定义"、"缩略语"。
        扫描整个章节（不因术语条目而中断），收集所有出现的子章节标题。
        """
        sub_sections = []
        if guide_idx < 0:
            return sub_sections
        expected = {"术语和定义", "缩略语"}
        for p in self.paragraphs[guide_idx + 1 :]:
            if p in expected and p not in sub_sections:
                sub_sections.append(p)
        return sub_sections

    def _looks_like_term_entry(self, paragraph: str) -> bool:
        """判断段落是否像术语条目（中英文混合或带冒号/括号，且非子章节标题）。"""
        # 如果是标准子章节标题，不算术语条目
        if paragraph in {"术语和定义", "缩略语"}:
            return False
        return bool(re.search(r"[：:]", paragraph)) or bool(
            re.search(r"[a-zA-Z]", paragraph) and re.search(r"[\u4e00-\u9fa5]", paragraph)
        )

    def _extract_acronym_section(self, guide_idx: int) -> List[str]:
        """提取"缩略语"子章节正文段落（用于识别已声明缩写）。"""
        if guide_idx < 0:
            return []
        capture = False
        result = []
        for p in self.paragraphs[guide_idx + 1 :]:
            if p == "缩略语":
                capture = True
                continue
            if capture:
                # 遇到其它子章节标题则停止
                if p in self.sub_sections and p != "缩略语":
                    break
                result.append(p)
        return result

    def _extract_declared_abbreviations(self, acronym_paras: List[str]) -> set:
        """从"缩略语"章节条目中提取已声明的缩写，如 'DRG：配电需求响应网关' -> {'DRG'}。"""
        declared = set()
        for p in acronym_paras:
            m = re.match(r"^([A-Za-z][A-Za-z0-9\-]*)\s*[：:]", p)
            if m:
                declared.add(m.group(1))
        return declared

    def _extract_body_paragraphs(self, guide_idx: int) -> List[str]:
        """提取章节正文段落（排除标题、引导语、子章节标题）。"""
        if guide_idx < 0:
            return []
        title = self.chapter_title if hasattr(self, "chapter_title") else ""
        body = []
        for i in range(guide_idx + 1, len(self.paragraphs)):
            p = self.paragraphs[i]
            if p == title or p in self.sub_sections:
                continue
            body.append(p)
        return body

    def _parse_entries(self, guide_idx: int) -> List[TermEntry]:
        """解析术语条目。"""
        entries: List[TermEntry] = []
        if guide_idx < 0:
            return entries

        current = TermEntry()
        for i in range(guide_idx + 1, len(self.paragraphs)):
            text = self.paragraphs[i]

            # 缩略语子章节之后为缩略语列表，不再是术语条目，停止解析
            if text == "缩略语":
                break
            # 其它子章节标题（如"术语和定义"）跳过，继续解析后续术语
            if text in self.sub_sections:
                continue

            # 术语行：按首行判断（支持名称与定义用换行分隔）
            if self._is_term_name_line(text):
                if current.name_cn or current.name_en or current.definition:
                    entries.append(current)
                current = TermEntry(raw_text=text, paragraph_index=i)
                first = text.split("\n")[0].strip()
                current.name_cn, current.name_en = self._split_term_name(first)
                # 首行之后（若有）作为定义
                rest = text[len(first):].strip("\n").strip()
                current.definition = rest
            else:
                current.definition += ("\n" if current.definition else "") + text

        if current.name_cn or current.name_en or current.definition:
            entries.append(current)

        return entries

    def _is_term_name_line(self, text: str) -> bool:
        """判断是否为术语名称行（按段落首行判断，支持名称与定义用换行分隔）。"""
        first = text.split("\n")[0].strip()
        if not first:
            return False
        # 含完整句末标点的首行不可能是术语名称行
        if first.endswith("。") or first.endswith("；") or first.endswith(";"):
            return False

        # 中英双语术语行：中文 + 空格/冒号 + 英文（允许任意大小写，由审核器校验小写）
        # 英文术语后可跟括号说明，如 (DR)、（OS-P）、【master station】
        if re.match(
            r"^[\u4e00-\u9fa5][\u4e00-\u9fa5\s、]*[\s:：]+[A-Za-z][A-Za-z0-9\s\-_/]*"
            + EN_TERM_BRACKET + r"$",
            first,
        ):
            return True

        # 中文术语（英文术语）
        if re.match(r"^[\u4e00-\u9fa5][\u4e00-\u9fa5\s、]*[（(][A-Za-z][A-Za-z0-9\s\-_/]*[）)]$", first):
            return True

        # 纯中文术语行（短，无定义内容）
        if re.match(r"^[\u4e00-\u9fa5\s、]+$", first) and len(first) <= 30:
            return True

        # 纯英文术语行（短）
        if re.match(r"^[A-Za-z][A-Za-z0-9\s\-_/]*$", first) and len(first) <= 40:
            return True

        return False

    def _split_term_name(self, text: str) -> Tuple[str, str]:
        """从术语行中分离中文术语和英文术语。"""
        # "中文术语 英文术语" 或 "中文术语：英文术语"（英文后可跟括号说明）
        m = re.match(
            r"^([\u4e00-\u9fa5][\u4e00-\u9fa5\s、]*)\s+([A-Za-z][A-Za-z0-9\s\-_/]*"
            + EN_TERM_BRACKET + r")$",
            text,
        )
        if m:
            return m.group(1).strip(), m.group(2).strip()

        # "中文术语（英文术语）"
        m = re.match(r"^([\u4e00-\u9fa5][\u4e00-\u9fa5\s、]*)[（(]([A-Za-z][A-Za-z0-9\s\-_/]*)[）)]$", text)
        if m:
            return m.group(1).strip(), m.group(2).strip()

        # 只有中文
        if re.match(r"^[\u4e00-\u9fa5\s、]+$", text):
            return text.strip(), ""

        # 只有英文
        if re.match(r"^[A-Za-z][A-Za-z0-9\s\-_/]*$", text):
            return "", text.strip()

        # 默认：整段作为中文
        return text, ""


# ---------------------------------------------------------------------------
# 审核器
# ---------------------------------------------------------------------------

class TerminologyAuditor:
    """根据规则对术语章节进行审核。"""

    def __init__(self, full_text: str = "", chapter: TerminologyChapter = None):
        self.full_text = full_text
        self.chapter = chapter or TerminologyChapter()
        self.issues: List[AuditIssue] = []

    # ------------------------------------------------------------------
    # 对外接口
    # ------------------------------------------------------------------

    def audit(self) -> List[AuditIssue]:
        """执行全部审核规则。"""
        self.issues = []
        self._check_chapter_title()
        self._check_guide_sentence_format()
        self._check_guide_sentence_location()
        self._check_guide_sentence_redundancy()
        self._check_no_terms_guide()
        self._check_sub_sections()
        self._check_entries_bilingual()
        self._check_entries_english_format()
        self._check_term_name_line()
        self._check_acronym_semicolon()
        self._check_definition_repeated_term()
        self._check_definition_requirement_words()
        self._check_definition_leading_words()
        self._check_terms_usage_in_body()
        return self.issues

    def to_json(self, indent: int = 2, ensure_ascii: bool = False) -> str:
        return json.dumps(
            [issue.to_dict() for issue in self.issues],
            indent=indent,
            ensure_ascii=ensure_ascii,
        )

    # ------------------------------------------------------------------
    # 规则实现
    # ------------------------------------------------------------------

    def _check_chapter_title(self):
        """检查章节名称是否规范。"""
        title = self.chapter.title
        if not title:
            self._add(
                rule="章节-术语和定义",
                location="章节标题",
                description="术语和定义章节结构错误",
                suggestion="检查章节是否缺失，格式是否错误。若没有术语和定义需要保留第3章并给出引导语：本文件没有术语和定义",
                paragraph_index=-2,
            )
            return

        if title not in TERMINOLOGY_CHAPTER_TITLES:
            # 若包含"术语"或"定义"字样但名称不规范
            if "术语" in title or "定义" in title:
                self._add(
                    rule="章节-术语和定义",
                    location=f"章节标题：{title}",
                    description="术语和定义章节名称描述错误",
                    suggestion="检查修改术语和定义章节名称",
                    paragraph_index=-2,
                )

    def _check_guide_sentence_format(self):
        """检查引导语格式是否为标准形式之一。"""
        guide = self.chapter.guide_sentence
        if not guide:
            return
        if not any(re.match(pattern, guide) for pattern in VALID_GUIDE_PATTERNS):
            self._add(
                rule="术语和定义-引导语",
                location=f"引导语：{guide}",
                description="引导语不符合要求。应满足以下三种情况之一：\n1. 下列术语和定义适用于本文件。\n2. ……界定的术语和定义适用于本文件。\n3. ……界定的以及下列术语和定义适用于本文件。\n如果没有术语和定义，应写：本文件没有需要界定的术语和定义。",
                suggestion="修改引导语为标准格式",
                paragraph_index=self.chapter.guide_paragraph_index,
            )

    def _check_guide_sentence_location(self):
        """检查是否能定位到引导语。"""
        if self.chapter.guide_paragraph_index < 0:
            self._add(
                rule="术语和定义-章节内容",
                location="术语章节",
                description="未找到引导语具体位置，请检查格式是否正确",
                suggestion="未找到引导语具体位置",
            )

    def _check_guide_sentence_redundancy(self):
        """检查引导语是否只有一句话（简单启发式：一个句号）。"""
        guide = self.chapter.guide_sentence
        if not guide:
            return
        sentence_endings = re.findall(r"。", guide)
        if len(sentence_endings) > 1 or len(guide) > 60:
            self._add(
                rule="术语和定义-章节内容",
                location=f"引导语：{guide}",
                description="引导语存在多余内容",
                suggestion="引导语只能有一句话，删除多余内容",
                paragraph_index=self.chapter.guide_paragraph_index,
            )

    def _check_no_terms_guide(self):
        """当无术语时，检查是否给出'本文件没有需要界定的术语和定义'。"""
        has_entries = bool(self.chapter.entries)
        guide = self.chapter.guide_sentence
        if not has_entries:
            if not re.match(r"^本文件没有需要界定的术语和定义。$", guide or ""):
                self._add(
                    rule="术语和定义-引导语",
                    location=f"引导语：{guide}",
                    description="如果没有术语和定义，应写：本文件没有需要界定的术语和定义。",
                    suggestion="修改引导语为标准格式",
                    paragraph_index=self.chapter.guide_paragraph_index,
                )

    def _check_sub_sections(self):
        """检查'术语、定义和缩略语'章节的子章节是否只包含两个小节。"""
        if self.chapter.title != "术语、定义和缩略语":
            return
        expected = {"术语和定义", "缩略语"}
        actual = set(self.chapter.sub_sections)
        if actual != expected:
            self._add(
                rule="术语和定义-章节内容",
                location=f"子章节：{self.chapter.sub_sections}",
                description="检查术语和定义章节是否真正包含缩略语。除了术语、定义和缩略语，不应有其他内容。",
                suggestion="检查修改术语和定义章节确保真正包含缩略语。除了术语、定义和缩略语，不应有其他内容。",
                paragraph_index=-2,
            )

    def _check_entries_bilingual(self):
        """检查术语条目是否同时包含中文和英文。"""
        for entry in self.chapter.entries:
            if not entry.name_cn and entry.name_en:
                self._add(
                    rule="术语和定义-术语和定义",
                    location=f"术语条目：{entry.raw_text}",
                    description="术语和定义应有对应的中文。",
                    suggestion="添加术语和定义应对应的中文。",
                    paragraph_index=entry.paragraph_index,
                )
            if entry.name_cn and not entry.name_en:
                self._add(
                    rule="术语和定义-术语和定义",
                    location=f"术语条目：{entry.raw_text}",
                    description="缺少术语和定义对应的英文单词翻译",
                    suggestion="添加术语和定义应对应的英文。",
                    paragraph_index=entry.paragraph_index,
                )

    def _check_entries_english_format(self):
        """检查英文术语是否全小写且仅含小写字母、空格和连字符。"""
        for entry in self.chapter.entries:
            en = entry.name_en
            if not en:
                continue
            # 忽略括号说明（如 (DR)、【master station】中的大写与括号）
            en_core = re.sub(r"[（(【][^）)】]*[）)】]", "", en)
            if re.search(r"[^a-z\s\-]", en_core):
                self._add(
                    rule="术语和定义-术语和定义",
                    location=f"术语条目：{entry.raw_text}",
                    description="术语和定义应有对应的英文，且英文必须是小写。",
                    suggestion="修改术语和定义应对应的英文为小写。",
                    paragraph_index=entry.paragraph_index,
                )

    def _check_term_name_line(self):
        """检查术语名称是否独占一行（不与定义在同一行）。

        仅在章节存在术语条目的前提下检查；无术语条目的章节（如引导语为
        "本文件没有需要界定的术语和定义。"的章节）其正文为叙述性内容，
        不适用该规则，直接跳过，避免误报。

        判定依据为段落首行：若首行本身即以句末标点结尾，说明术语名与定义
        混在同一行（未换行）；若术语名独占首行、定义在后续行（即使同属一个
        段落/含换行），则不视为违规。
        """
        if not self.chapter.entries:
            return
        src = self.chapter.source_paragraphs
        for p in self.chapter.body_paragraphs or []:
            first_line = p.split("\n")[0].strip()
            # 仅当首行（术语名所在行）本身含句末标点，才说明未换行
            if not first_line.endswith("。"):
                continue
            # 首行符合术语名模式且带定义句号
            if len(first_line) > 30 and self._prefix_looks_like_term_name(first_line):
                idx = src.index(p) if p in src else -1
                self._add(
                    rule="术语和定义-章节内容",
                    location=f"术语条目：{p}",
                    description="术语未换行",
                    suggestion="将术语换行",
                    paragraph_index=idx,
                )

    def _prefix_looks_like_term_name(self, text: str) -> bool:
        """判断文本前缀是否符合术语名称模式（必须以术语名开头）。"""
        # 移除可能存在的括号说明后再判断
        text = re.sub(r"[（(].*?[）)]", "", text).strip()
        if not text:
            return False
        # 必须以中文术语开头，后接英文（"中文 英文" 或 "中文（英文"）
        if re.match(r"^[\u4e00-\u9fa5][\u4e00-\u9fa5\s、]*[\s:：]+[A-Za-z]", text):
            return True
        if re.match(r"^[\u4e00-\u9fa5][\u4e00-\u9fa5\s、]*[（(][A-Za-z]", text):
            return True
        return False

    def _check_acronym_semicolon(self):
        """检查术语定义中包含括号说明的缩略语时，括号前是否用中文分号。

        已在"缩略语"子章节声明过的缩写属于交叉引用，不在此检查范围，予以排除。
        """
        declared = getattr(self.chapter, "declared_abbreviations", set())
        for entry in self.chapter.entries:
            definition = entry.definition
            for m in re.finditer(r"[（(]([A-Za-z][A-Za-z0-9\-]*)[）)]", definition):
                abbr = m.group(1)
                # 已声明的缩写属于交叉引用，不要求使用分号分隔
                if abbr in declared:
                    continue
                before = definition[m.start() - 1] if m.start() > 0 else ""
                if before != "；":
                    self._add(
                        rule="术语和定义-章节内容",
                        location=f"术语'{entry.name_cn or entry.raw_text}'定义",
                        description="术语缩略语未用；分割",
                        suggestion="在缩略语括号前使用中文分号'；'分隔",
                        paragraph_index=entry.paragraph_index,
                    )
                    break

    def _check_definition_repeated_term(self):
        """检查定义正文是否重复出现具体术语名称。

        排除"注"以及"[来源：…]""【来源：…】"等出处标注——其中可能包含术语名
        （如"电力需求响应系统通用技术规范"含有"需求响应"），不应误判为定义重复。
        """
        for entry in self.chapter.entries:
            definition = entry.definition
            if not definition:
                continue
            # 移除半角/方头括号的出处标注及"注"
            cleaned = re.sub(r"\[[^\]]*\]", "", definition)
            cleaned = re.sub(r"【[^】]*】", "", cleaned)
            cleaned = re.sub(r"注[：:].*", "", cleaned, flags=re.S)
            if entry.name_cn and entry.name_cn in cleaned:
                self._add(
                    rule="术语和定义-后续使用",
                    location=f"术语'{entry.name_cn}'定义",
                    description="具体的术语不应在定义中重复出现",
                    suggestion="检查术语使用规范性",
                    paragraph_index=entry.paragraph_index,
                )

    def _check_definition_requirement_words(self):
        """检查术语定义中是否包含要求性条款词汇。

        先剔除含"应/可/宜"的非要求性复合词（响应、应用、可能、可行…），
        避免单字"应/可"误命中。
        """
        pattern = re.compile("(" + "|".join(map(re.escape, REQUIREMENT_WORDS)) + ")")
        for entry in self.chapter.entries:
            definition = entry.definition
            cleaned = definition
            for c in SAFE_COMPOUNDS_WITH_MODAL:
                cleaned = cleaned.replace(c, "")
            matches = pattern.findall(cleaned)
            if matches:
                self._add(
                    rule="术语和定义-条款约束",
                    location=f"术语'{entry.name_cn or entry.raw_text}'定义",
                    description="术语和定义中不应包含要求性条款（应、宜、可、应该、建议、可以、不应、不准许、无须）",
                    suggestion="检查修改术语和定义",
                    details={"命中词汇": list(set(matches))},
                    paragraph_index=entry.paragraph_index,
                )

    def _check_definition_leading_words(self):
        """检查定义是否以'系指/是指/指/是'等引导词开头。"""
        for entry in self.chapter.entries:
            definition = entry.definition
            if not definition:
                continue
            first_sentence = definition.split("。")[0]
            for word in FORBIDDEN_LEADING_WORDS:
                if first_sentence.startswith(word):
                    self._add(
                        rule="术语和定义-条款约束",
                        location=f"术语'{entry.name_cn or entry.raw_text}'定义",
                        description="定义中不应使用'是指', '指', '是'等字，请改正",
                        suggestion="术语和定义中不应使用（指，是，是指，系指）",
                        paragraph_index=entry.paragraph_index,
                    )
                    break

    def _check_terms_usage_in_body(self):
        """统计全文中每个中文术语出现次数，仅在术语条目中出现则建议删除。"""
        if not self.full_text or not self.chapter.entries:
            return
        for entry in self.chapter.entries:
            cn = entry.name_cn
            if not cn or len(cn) < 2:
                continue
            count = len(re.findall(re.escape(cn), self.full_text))
            if count <= 1:
                self._add(
                    rule="术语和定义-正文应用",
                    location=f"术语'{cn}'",
                    description="该术语、定义未在正文中被使用，建议删除",
                    suggestion="删除该术语和定义",
                    severity="warning",
                    details={"出现次数": count},
                    paragraph_index=entry.paragraph_index,
                )

    # ------------------------------------------------------------------
    # 工具方法
    # ------------------------------------------------------------------

    def _add(self, rule="", location="", description="", suggestion="",
             severity="error", details=None, paragraph_index=-1):
        self.issues.append(
            AuditIssue(
                rule=rule,
                location=location,
                description=description,
                suggestion=suggestion,
                severity=severity,
                details=details or {},
                paragraph_index=paragraph_index,
            )
        )


# ---------------------------------------------------------------------------
# 便捷入口
# ---------------------------------------------------------------------------

def audit_terminology(
    paragraphs: List[str],
    chapter_title: str = "",
    full_text: str = "",
) -> List[AuditIssue]:
    """
    对术语章节段落进行审核。

    参数:
        paragraphs: 术语章节内的段落文本列表。
        chapter_title: 章节标题；为空时自动检测。
        full_text: 完整文档文本，用于统计术语在正文中的使用次数。
    """
    parser = TerminologyParser(paragraphs)
    chapter = parser.parse(chapter_title)
    auditor = TerminologyAuditor(full_text=full_text, chapter=chapter)
    return auditor.audit()


def audit_docx(docx_path: str) -> Tuple[str, "TerminologyChapter", List[AuditIssue], List, object, object]:
    """
    对 .docx 文件中的术语章节进行审核。

    返回: (章节标题, 章节对象, 问题列表, 章节段落对应的 Paragraph 对象列表,
           章节标题对应的 Paragraph 对象, 已加载的 Document 对象)
        - 后三项用于把问题回写为 Word 批注时定位锚点。
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("请安装 python-docx: pip install python-docx") from exc

    doc = Document(docx_path)
    # 仅保留非空段落，记录 (段落对象, 文本)
    items = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            items.append((p, t))

    # 定位术语章节
    chapter_idx = -1
    chapter_title = ""
    for i, (p, t) in enumerate(items):
        if t in TERMINOLOGY_CHAPTER_TITLES:
            chapter_idx = i
            chapter_title = t
            break

    if chapter_idx < 0:
        return "", TerminologyChapter(), [
            AuditIssue(
                rule="章节-术语和定义",
                location="文档",
                description="术语和定义章节结构错误",
                suggestion="检查章节是否缺失，格式是否错误。若没有术语和定义需要保留第3章并给出引导语：本文件没有术语和定义",
            )
        ], [], None, doc

    title_para = items[chapter_idx][0]

    # 截取术语章节段落（直到下一个同层级章节或文档结束）
    # 允许存在的术语章节子标题集合（不视为章节终止）
    allowed_in_chapter = {chapter_title, "缩略语"}
    chapter_span = []  # [(段落对象, 文本)]
    for i in range(chapter_idx + 1, len(items)):
        p, t = items[i]
        if _is_top_level_heading(p.style.name, t, allowed_in_chapter):
            break
        chapter_span.append((p, t))

    chapter_paragraphs = [t for _, t in chapter_span]
    para_objects = [p for p, _ in chapter_span]

    full_text = "\n".join(t for _, t in items)
    parser = TerminologyParser(chapter_paragraphs)
    chapter = parser.parse(chapter_title)
    auditor = TerminologyAuditor(full_text=full_text, chapter=chapter)
    issues = auditor.audit()
    return chapter_title, chapter, issues, para_objects, title_para, doc


def _is_top_level_heading(style: str, text: str, allowed_in_chapter: set) -> bool:
    """判断段落是否为顶层章节标题（术语章节之后的下一个章节）。

    以段落样式为主要依据（如含"章标题"/Heading 等），并兼容无样式时的
    编号章节（如 "4 总体要求"）作为兜底。
    """
    style_l = (style or "").lower()
    is_heading_style = (
        "章标题" in (style or "")
        or style_l == "heading 1"
        or "1级标题" in (style or "")
        or "标题 1" in (style or "")
    )
    if is_heading_style:
        # 属于术语章节自身的子标题（如"缩略语"）则不视为终止
        return text not in allowed_in_chapter
    # 兜底：编号章节标题，如 "4 总体要求" / "五、xxx"
    return bool(re.match(r"^[\d一二三四五六七八九十]+[\.\s、]", text)) and len(text) <= 30


# ---------------------------------------------------------------------------
# Word 批注回写
# ---------------------------------------------------------------------------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"


def _get_or_create_comments_part(doc):
    """获取或新建 comments.xml 部件（Word 批注容器）。"""
    for rel in doc.part.rels.values():
        if rel.reltype == COMMENTS_REL:
            return rel.target_part
    element = OxmlElement("w:comments")
    partname = PackURI("/word/comments.xml")
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
    )
    part = XmlPart(partname, content_type, element, doc.part.package)
    doc.part.relate_to(part, COMMENTS_REL)
    return part


def _add_comment_def(comments_part, cid: int, author: str, lines: List[str]):
    """在 comments.xml 中追加一条批注定义。"""
    root = comments_part.element
    comment = OxmlElement("w:comment")
    comment.set(qn("w:id"), str(cid))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:date"), datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"))
    comment.set(qn("w:initials"), author[:1] or "审")
    p = OxmlElement("w:p")
    for idx, line in enumerate(lines):
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r.append(t)
        if idx < len(lines) - 1:
            r.append(OxmlElement("w:br"))
        p.append(r)
    comment.append(p)
    root.append(comment)


def _anchor_comment(paragraph, cid: int):
    """在段落中插入批注范围标记与引用，使批注锚定到该段落。"""
    p = paragraph._p
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(cid))
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        p.insert(0, start)
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(cid))
    ref = OxmlElement("w:r")
    ref.append(OxmlElement("w:commentReference"))
    ref[0].set(qn("w:id"), str(cid))
    p.append(end)
    p.append(ref)


def _resolve_anchor(issue, para_objects, title_para):
    """根据问题的 paragraph_index 解析出要锚定的段落对象。"""
    pi = issue.paragraph_index
    if pi == -2:
        return title_para
    if pi is not None and pi >= 0 and pi < len(para_objects):
        return para_objects[pi]
    return None


def _build_comment_lines(issue: AuditIssue) -> List[str]:
    """构造一条批注的正文行。

    所有规则的 docx 批注正文都不展示「【规则名】[级别]」标签前缀，
    仅保留「问题 / 建议 / 详情」正文，使批注更贴近人工修订意见、不显突兀。
    """
    lines = [
        f"问题：{issue.description}",
        f"建议：{issue.suggestion}",
    ]
    if issue.details:
        extra = "；".join(f"{k}={v}" for k, v in issue.details.items())
        lines.append(f"详情：{extra}")
    return lines


def write_annotated_docx(
    doc,
    issues: List[AuditIssue],
    para_objects: List,
    title_para,
    out_path: str,
    author: str = "术语审核",
):
    """
    把审核问题作为 Word 批注（comments）写回源文档，保存为 out_path。

    doc 必须是与 para_objects / title_para 同源的、已加载的 Document 对象。
    """
    comments_part = _get_or_create_comments_part(doc)
    existing = [
        int(c.get(qn("w:id")))
        for c in comments_part.element.findall(qn("w:comment"))
        if c.get(qn("w:id"), "").isdigit()
    ]
    next_id = (max(existing) + 1) if existing else 0

    for issue in issues:
        lines = _build_comment_lines(issue)
        _add_comment_def(comments_part, next_id, author, lines)
        anchor = _resolve_anchor(issue, para_objects, title_para)
        if anchor is not None:
            _anchor_comment(anchor, next_id)
        next_id += 1

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# xlsx 汇总导出
# ---------------------------------------------------------------------------

def write_issues_xlsx(
    issues: List[AuditIssue],
    out_path: str,
    meta: Optional[dict] = None,
):
    """把审核问题导出为 xlsx 汇总表（含概览与问题明细两个工作表）。"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    # ---- 概览表 ----
    ws_summary = wb.active
    ws_summary.title = "概览"
    meta = meta or {}
    summary_rows = [
        ("章节标题", meta.get("章节标题", "")),
        ("引导语", meta.get("引导语", "")),
        ("术语条目数", meta.get("术语条目数", "")),
        ("子章节", "、".join(meta.get("子章节", []) or [])),
        ("问题总数", meta.get("问题总数", len(issues))),
        ("审核结论", meta.get("审核结论", "")),
    ]
    ws_summary.append(["项目", "内容"])
    for k, v in summary_rows:
        ws_summary.append([k, str(v)])
    hdr_fill = PatternFill("solid", fgColor="1F4E78")
    hdr_font = Font(bold=True, color="FFFFFF")
    ws_summary["A1"].fill = hdr_fill
    ws_summary["A1"].font = hdr_font
    ws_summary["B1"].fill = hdr_fill
    ws_summary["B1"].font = hdr_font
    ws_summary.column_dimensions["A"].width = 14
    ws_summary.column_dimensions["B"].width = 80
    for row in ws_summary.iter_rows(min_row=2):
        row[1].alignment = Alignment(wrap_text=True, vertical="top")

    # ---- 问题明细表 ----
    ws = wb.create_sheet("审核问题")
    headers = ["序号", "审核要点", "级别", "问题描述", "修改建议", "位置", "详情", "段落索引"]
    ws.append(headers)
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=c)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    sev_fill = {
        "error": PatternFill("solid", fgColor="F8CBAD"),
        "warning": PatternFill("solid", fgColor="FFE699"),
        "info": PatternFill("solid", fgColor="DDEBF7"),
    }
    for i, iss in enumerate(issues, 1):
        detail = "；".join(f"{k}={v}" for k, v in iss.details.items()) if iss.details else ""
        ws.append([
            i, iss.rule, iss.severity, iss.description,
            iss.suggestion, iss.location, detail, iss.paragraph_index,
        ])
        r = ws.max_row
        fill = sev_fill.get(iss.severity)
        if fill:
            ws.cell(row=r, column=3).fill = fill
        for c in range(1, len(headers) + 1):
            ws.cell(row=r, column=c).border = border
            ws.cell(row=r, column=c).alignment = Alignment(
                wrap_text=True, vertical="top"
            )

    widths = [6, 22, 8, 46, 40, 34, 30, 10]
    for idx, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = w
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws.max_row}"

    wb.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# 命令行入口
# ---------------------------------------------------------------------------

def main():
    import argparse

    parser = argparse.ArgumentParser(description="术语章节审核器（含批注回写与 xlsx 汇总）")
    parser.add_argument("input", help="输入 .docx 文件或 JSON 文件（段落列表）")
    parser.add_argument("-o", "--output", help="输出 JSON 结果文件路径（可选）")
    parser.add_argument("--docx-out", help="带批注的 docx 输出路径（仅 .docx 输入；默认 <输入名>_术语审核批注.docx）")
    parser.add_argument("--xlsx-out", help="审核问题汇总 xlsx 输出路径（默认 <输入名>_术语审核问题.xlsx）")
    parser.add_argument(
        "--title",
        default="",
        help="章节标题，为空时自动检测",
    )
    parser.add_argument(
        "--no-comments",
        action="store_true",
        help="仅输出 xlsx，不向 docx 写批注",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误：文件不存在 {input_path}")
        return 1

    is_docx = input_path.suffix.lower() == ".docx"

    if is_docx:
        title, chapter, issues, para_objects, title_para, doc = audit_docx(str(input_path))
    else:
        with open(input_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            paragraphs = data.get("paragraphs", [])
            full_text = data.get("full_text", "")
        else:
            paragraphs = data
            full_text = ""
        title = args.title
        ch_parser = TerminologyParser(paragraphs)
        chapter = ch_parser.parse(title)
        title = chapter.title or title or "术语章节"
        auditor = TerminologyAuditor(full_text=full_text, chapter=chapter)
        issues = auditor.audit()
        para_objects, title_para, doc = [], None, None

    result = {
        "章节标题": title,
        "引导语": chapter.guide_sentence,
        "术语条目数": len(chapter.entries),
        "子章节": chapter.sub_sections,
        "问题总数": len(issues),
        "审核结论": "通过" if not issues else "存在问题",
        "问题列表": [issue.to_dict() for issue in issues],
    }

    # 输出 JSON（可选）
    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(json.dumps(result, ensure_ascii=False, indent=2))
        print(f"JSON 结果已保存：{args.output}")

    # 默认输出路径（基于输入文件名）
    if args.xlsx_out:
        xlsx_path = args.xlsx_out
    else:
        xlsx_path = str(input_path.with_name(input_path.stem + "_术语审核问题.xlsx"))
    write_issues_xlsx(issues, xlsx_path, meta=result)
    print(f"xlsx 汇总已保存：{xlsx_path}")

    # 向 docx 写批注（仅 .docx 输入且未禁用）
    if is_docx and not args.no_comments:
        if args.docx_out:
            docx_out = args.docx_out
        else:
            docx_out = str(input_path.with_name(input_path.stem + "_术语审核批注.docx"))
        if doc is not None:
            write_annotated_docx(doc, issues, para_objects, title_para, docx_out)
            print(f"带批注 docx 已保存：{docx_out}")
        else:
            print("警告：无法写批注（Document 对象缺失）。")

    print(f"审核完成：共 {len(issues)} 个问题。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
