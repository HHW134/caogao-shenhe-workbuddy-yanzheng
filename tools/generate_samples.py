#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成用于端到端验证的样例 docx(含有意违规)。

内容为合成的虚构标准文本，不含任何真实草案，可安全放入仓库。
默认输出到 <仓库>/samples，可用第一个命令行参数指定其它目录：

    python tools/generate_samples.py [输出目录]
"""
import os
import sys

from docx import Document

SAMPLES = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "samples")
if len(sys.argv) > 1:
    SAMPLES = sys.argv[1]


def sample1_normal_violations():
    """常规段落条目, 覆盖 A4R003~A4R015(除 A4R002)。"""
    d = Document()
    d.add_heading("XXXX 标准", level=0)
    d.add_heading("1 范围", level=1)
    d.add_paragraph("本文件规定了……")
    d.add_heading("2 规范性引用文件", level=1)
    # 引导语(正确)
    d.add_paragraph(
        "下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款。"
        "其中，注日期的引用文件，仅该日期对应的版本适用于本文件；"
        "不注日期的引用文件，其最新版本（包括所有的修改单）适用于本文件。"
    )
    # 条目(故意乱序/含各种违规)
    d.add_paragraph("GB/T 12345—2020 信息技术 通用规范")          # 正确, 正文规范引用
    d.add_paragraph("YD/T 987—2019 某通信行业标准")               # 正确, 正文规范引用
    d.add_paragraph("ISO/IEC 2382:2015 信息技术 词汇")            # 国际缺英文(A4R006)
    d.add_paragraph("GB/T 54321-2020 某标准")                     # 连字符年份(A4R008)
    d.add_paragraph("3、GB/T 11111—2020 某标准")                  # 人工序号(A4R009)
    d.add_paragraph("《GB/T 22222—2020 某标准》")                 # 书名号(A4R009), 正文未引用(A4R010)
    d.add_paragraph("ISO/IEC 27001:2013 信息安全")                # 国际缺英文(A4R006)
    d.add_paragraph("T/ABC 123—2020 团体标准规范")                # 团体, 正文规范引用
    d.add_paragraph("DB 11/T 456—2018 地方标准规范")              # 地方, 正文规范引用
    d.add_paragraph("GB/T XXXX—2020 占位符标准")                  # 占位符(A4R015), 未引用(A4R010)
    d.add_paragraph("某规范名称 没有编号的标准")                    # 缺编号(A4R005), 未引用(A4R010)
    d.add_paragraph("H.323 多媒体通信协议")                       # 缺组织代号(A4R007), 未引用(A4R010)
    d.add_paragraph("RFC 2119 工作要求")                          # 仅在注中提及 -> A4R011
    d.add_paragraph("IEEE 802.3 以太网规范")                      # 仅在来源中提及 -> A4R012

    d.add_heading("3 术语和定义", level=1)
    d.add_paragraph("下列术语和定义适用于本文件。")
    d.add_paragraph("3.1 术语A")
    d.add_paragraph("来源：IEEE 802.3")                           # A4R012 来源

    d.add_heading("4 要求", level=1)
    d.add_paragraph("产品应符合 GB/T 12345—2020 的规定。")          # 规范引用, 正确
    d.add_paragraph("系统应按照 YD/T 987—2019 的要求设计。")         # 规范引用, 正确
    d.add_paragraph("接口应满足 T/ABC 123—2020 的约定。")            # 规范引用, 正确
    d.add_paragraph("网络应依据 DB 11/T 456—2018 实施。")            # 规范引用, 正确
    d.add_paragraph("术语应符合 ISO/IEC 2382:2015 的定义。")          # 国际, 条目缺英文(A4R006)
    d.add_paragraph("安全应符合 ISO/IEC 27001:2013 的要求。")        # 国际, 条目缺英文(A4R006)
    d.add_paragraph("注：本文件部分内容参考 RFC 2119。")            # 仅在注 -> A4R011 (非规范引用)
    d.add_paragraph("产品应符合 GB/T 12345—2010 的要求。")          # 年代号不一致 -> A4R013
    d.add_paragraph("测试方法应按照 GB/T 99999 的规定执行。")        # 正文引用但不在第二章 -> A4R014

    out = os.path.join(SAMPLES, "sample1_violations.docx")
    d.save(out)
    print("saved", out)


def sample2_table_form():
    """章节内容以表格呈现 -> A4R002。"""
    d = Document()
    d.add_heading("YYYY 标准", level=0)
    d.add_heading("1 范围", level=1)
    d.add_paragraph("本文件规定了……")
    d.add_heading("2 规范性引用文件", level=1)
    d.add_paragraph("下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款。")
    table = d.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "GB/T 12345—2020 信息技术"
    table.cell(1, 0).text = "YD/T 987—2019 某行业标准"
    table.cell(2, 0).text = "ISO/IEC 2382:2015 信息技术 词汇"
    d.add_heading("3 要求", level=1)
    d.add_paragraph("产品应符合 GB/T 12345—2020 的规定。")
    out = os.path.join(SAMPLES, "sample2_table.docx")
    d.save(out)
    print("saved", out)


def sample3_missing_chapter():
    """缺少规范性引用文件章节 -> A4R001。"""
    d = Document()
    d.add_heading("ZZZZ 标准", level=0)
    d.add_heading("1 范围", level=1)
    d.add_paragraph("本文件规定了……")
    d.add_heading("2 术语和定义", level=1)
    d.add_paragraph("下列术语和定义适用于本文件。")
    d.add_heading("3 要求", level=1)
    d.add_paragraph("产品应符合 GB/T 12345—2020 的规定。")
    out = os.path.join(SAMPLES, "sample3_missing.docx")
    d.save(out)
    print("saved", out)


def sample4_missing_guide():
    """有引用条目但缺少标准引导语 -> A4R003。"""
    d = Document()
    d.add_heading("WWWW 标准", level=0)
    d.add_heading("1 范围", level=1)
    d.add_paragraph("本文件规定了……")
    d.add_heading("2 规范性引用文件", level=1)
    # 故意不写引导语, 直接列条目
    d.add_paragraph("GB/T 12345—2020 信息技术 通用规范")
    d.add_paragraph("YD/T 987—2019 某通信行业标准")
    d.add_heading("3 要求", level=1)
    d.add_paragraph("产品应符合 GB/T 12345—2020 的规定。")
    d.add_paragraph("系统应按照 YD/T 987—2019 的要求设计。")
    out = os.path.join(SAMPLES, "sample4_missing_guide.docx")
    d.save(out)
    print("saved", out)


def sample5_compliant():
    """完全合规: 引导语/排序/编号/格式/年份/正文引用 均正确。"""
    d = Document()
    d.add_heading("合规 标准", level=0)
    d.add_heading("1 范围", level=1)
    d.add_paragraph("本文件规定了……")
    d.add_heading("2 规范性引用文件", level=1)
    d.add_paragraph(
        "下列文件中的内容通过文中的规范性引用而构成本文件必不可少的条款。"
        "其中，注日期的引用文件，仅该日期对应的版本适用于本文件；"
        "不注日期的引用文件，其最新版本（包括所有的修改单）适用于本文件。"
    )
    d.add_paragraph("GB/T 12345—2020 信息技术 通用规范")
    d.add_paragraph("GB/T 22222—2020 某通用规范")
    d.add_paragraph("YD/T 987—2019 某通信行业标准")
    d.add_paragraph("DB 11/T 456—2018 地方标准规范")
    d.add_paragraph("T/ABC 123—2020 团体标准规范")
    d.add_paragraph("IEEE 802.3 以太网规范（Ethernet）")
    d.add_paragraph("ISO/IEC 2382:2015 信息技术 词汇（Information technology — Vocabulary）")
    d.add_heading("3 要求", level=1)
    d.add_paragraph("产品应符合 GB/T 12345—2020 的规定。")
    d.add_paragraph("系统应符合 GB/T 22222—2020 的要求。")
    d.add_paragraph("接口应按照 YD/T 987—2019 设计。")
    d.add_paragraph("网络应依据 DB 11/T 456—2018 实施。")
    d.add_paragraph("协议应满足 T/ABC 123—2020 的约定。")
    d.add_paragraph("术语应符合 ISO/IEC 2382:2015 的定义。")
    d.add_paragraph("安全应符合 IEEE 802.3 的要求。")
    out = os.path.join(SAMPLES, "sample5_compliant.docx")
    d.save(out)
    print("saved", out)


def sample6_no_refs():
    """无引用文件, 仅保留引导语, 应合规(0 问题)。"""
    d = Document()
    d.add_heading("无引用 标准", level=0)
    d.add_heading("1 范围", level=1)
    d.add_paragraph("本文件规定了……")
    d.add_heading("2 规范性引用文件", level=1)
    d.add_paragraph("本文件没有规范性引用文件。")
    d.add_heading("3 要求", level=1)
    d.add_paragraph("产品应满足下列要求。")
    out = os.path.join(SAMPLES, "sample6_no_refs.docx")
    d.save(out)
    print("saved", out)


if __name__ == "__main__":
    os.makedirs(SAMPLES, exist_ok=True)
    sample1_normal_violations()
    sample2_table_form()
    sample3_missing_chapter()
    sample4_missing_guide()
    sample5_compliant()
    sample6_no_refs()
